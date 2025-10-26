"""
Interactive Mode for CLI - FIXED VERSION with File Selection
"""

import cmd
import sys
import os
import shlex
from pathlib import Path
import glob
from typing import Dict, Any  # ADD THIS IMPORT

class InteractiveMode(cmd.Cmd):
    """
    Interactive command shell for AI File Converter
    With file selection and conversion functionality
    """
    
    def __init__(self):
        super().__init__()
        self.prompt = "🤖 converter> "
        self.intro = """
╔══════════════════════════════════════════════╗
║           AI FILE CONVERTER SHELL            ║
║          Interactive File Selection          ║
╚══════════════════════════════════════════════╝
Type 'help' for commands, 'scan' to find files, 'convert' to convert
"""
        self.current_directory = os.getcwd()
        self.selected_files = []
        
    def do_help(self, arg):
        """Show available commands"""
        commands = {
            "help": "Show this help message",
            "exit": "Exit the interactive shell",
            "list": "List available conversion tools",
            "status": "Show system status",
            "scan": "Scan directory for convertible files",
            "select": "Select files for conversion - usage: select <pattern> or select <file1> <file2>...",
            "selected": "Show currently selected files",
            "clear_selection": "Clear file selection",
            "convert": "Convert selected files - usage: convert <operation> [output_dir]",
            "cd": "Change directory",
            "pwd": "Show current directory",
            "models": "Show available AI models",
            "clear": "Clear the screen"
        }
        
        print("\n📋 AVAILABLE COMMANDS:")
        print("=" * 50)
        for cmd, desc in commands.items():
            print(f"  {cmd:20} - {desc}")
        print()

    def do_models(self, arg):
        """Show available AI models"""
        try:
            # Import the fixed model manager
            from ai_infrastructure.ollama.model_manager import model_manager
            
            print("🤖 Checking available AI models...")
            
            # Get available models
            models = model_manager.get_available_models()
            
            if models:
                print(f"\n🎯 AVAILABLE AI MODELS ({len(models)}):")
                print("=" * 60)
                for model in models:
                    model_name = model.get('name', 'Unknown')
                    model_size = model.get('size', 'Unknown size')
                    modified = model.get('modified_at', 'Unknown date')
                    
                    print(f"📚 {model_name}")
                    print(f"   📏 Size: {model_size}")
                    print(f"   📅 Modified: {modified}")
                    print("-" * 40)
            else:
                print("❌ No models found or Ollama is not running")
                print("💡 Make sure:")
                print("   1. Ollama is running: ollama serve")
                print("   2. You have models pulled: ollama pull mistral")
                print("   3. Check models: ollama list")
                
        except ImportError as e:
            print(f"❌ Model manager not available: {e}")
            print("💡 Check if ai_infrastructure.ollama.model_manager exists")
        except Exception as e:
            print(f"❌ Error checking models: {e}")
    
    def do_cd(self, arg):
        """Change current directory"""
        if not arg:
            print(f"Current directory: {self.current_directory}")
            return
        
        new_dir = arg.strip()
        if not os.path.exists(new_dir):
            print(f"❌ Directory not found: {new_dir}")
            return
        
        if not os.path.isdir(new_dir):
            print(f"❌ Not a directory: {new_dir}")
            return
        
        try:
            os.chdir(new_dir)
            self.current_directory = os.getcwd()
            print(f"📁 Changed to directory: {self.current_directory}")
        except Exception as e:
            print(f"❌ Failed to change directory: {e}")
    
    def do_pwd(self, arg):
        """Show current working directory"""
        print(f"📁 Current directory: {self.current_directory}")
    
    def do_scan(self, arg):
        """Scan directory for convertible files"""
        directory = arg if arg else self.current_directory
        
        if not os.path.exists(directory):
            print(f"❌ Directory not found: {directory}")
            return
        
        if not os.path.isdir(directory):
            print(f"❌ Not a directory: {directory}")
            return
        
        print(f"🔍 Scanning: {directory}")
        
        # Define supported file types for conversion
        file_patterns = [
            "*.pdf", "*.docx", "*.doc", "*.csv", 
            "*.txt", "*.jpg", "*.jpeg", "*.png", "*.gif"
        ]
        
        convertible_files = []
        for pattern in file_patterns:
            files = glob.glob(os.path.join(directory, pattern))
            convertible_files.extend(files)
        
        # Also search recursively
        for root, dirs, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                if any(file.lower().endswith(ext) for ext in ['.pdf', '.docx', '.doc', '.csv', '.txt', '.jpg', '.jpeg', '.png', '.gif']):
                    convertible_files.append(file_path)
        
        # Remove duplicates and sort
        convertible_files = sorted(list(set(convertible_files)))
        
        if convertible_files:
            print(f"\n📁 Found {len(convertible_files)} convertible files:")
            for i, file_path in enumerate(convertible_files[:20], 1):  # Show first 20
                file_size = os.path.getsize(file_path)
                size_str = self._format_file_size(file_size)
                print(f"  {i:2d}. {os.path.basename(file_path):30} ({size_str:>8}) - {file_path}")
            
            if len(convertible_files) > 20:
                print(f"   ... and {len(convertible_files) - 20} more files")
            
            print(f"\n💡 Use 'select <pattern>' or 'select <number>' to choose files")
        else:
            print("❌ No convertible files found")
            print("   Supported formats: PDF, DOCX, DOC, CSV, TXT, JPG, JPEG, PNG, GIF")
    
    def do_select(self, arg):
        """Select files for conversion"""
        if not arg:
            print("❌ Usage: select <pattern> or select <number> or select <file1> <file2>...")
            print("   Examples:")
            print("     select *.pdf              - Select all PDF files")
            print("     select 1 3 5              - Select files by number from scan results")
            print("     select file1.pdf file2.docx - Select specific files")
            return
        
        args = shlex.split(arg)
        
        # Check if args are numbers (from scan results)
        if all(a.isdigit() for a in args):
            self._select_by_numbers(args)
        # Check if args are file patterns
        elif any('*' in a or '?' in a for a in args):
            self._select_by_pattern(args)
        # Assume args are specific file paths
        else:
            self._select_by_paths(args)
    
    def _select_by_numbers(self, numbers):
        """Select files by numbers from scan results"""
        # We need scan results - for simplicity, scan current directory
        directory = self.current_directory
        convertible_files = self._scan_directory(directory)
        
        selected = []
        for num_str in numbers:
            try:
                num = int(num_str)
                if 1 <= num <= len(convertible_files):
                    file_path = convertible_files[num-1]
                    if os.path.exists(file_path):
                        selected.append(file_path)
                    else:
                        print(f"❌ File not found: {file_path}")
                else:
                    print(f"❌ Invalid number: {num}. Available: 1-{len(convertible_files)}")
            except ValueError:
                print(f"❌ Invalid number: {num_str}")
        
        self.selected_files.extend(selected)
        print(f"✅ Selected {len(selected)} files")
        self.do_selected(None)
    
    def _select_by_pattern(self, patterns):
        """Select files by pattern matching"""
        selected = []
        for pattern in patterns:
            files = glob.glob(os.path.join(self.current_directory, pattern))
            if files:
                selected.extend(files)
            else:
                print(f"❌ No files match pattern: {pattern}")
        
        self.selected_files.extend(selected)
        print(f"✅ Selected {len(selected)} files matching patterns")
        self.do_selected(None)
    
    def _select_by_paths(self, paths):
        """Select files by specific paths"""
        selected = []
        for path in paths:
            full_path = os.path.join(self.current_directory, path) if not os.path.isabs(path) else path
            if os.path.exists(full_path):
                selected.append(full_path)
            else:
                print(f"❌ File not found: {path}")
        
        self.selected_files.extend(selected)
        print(f"✅ Selected {len(selected)} files")
        self.do_selected(None)
    
    def _scan_directory(self, directory):
        """Scan directory and return list of convertible files"""
        file_patterns = ["*.pdf", "*.docx", "*.doc", "*.csv", "*.txt", "*.jpg", "*.jpeg", "*.png", "*.gif"]
        convertible_files = []
        
        for pattern in file_patterns:
            files = glob.glob(os.path.join(directory, pattern))
            convertible_files.extend(files)
        
        return sorted(list(set(convertible_files)))
    
    def do_selected(self, arg):
        """Show currently selected files"""
        if not self.selected_files:
            print("📭 No files selected")
            print("💡 Use 'scan' to find files and 'select' to choose them")
            return
        
        print(f"\n📋 Currently selected files ({len(self.selected_files)}):")
        print("=" * 60)
        total_size = 0
        for i, file_path in enumerate(self.selected_files, 1):
            if os.path.exists(file_path):
                file_size = os.path.getsize(file_path)
                total_size += file_size
                size_str = self._format_file_size(file_size)
                print(f"  {i:2d}. {os.path.basename(file_path):30} ({size_str:>8}) - {file_path}")
            else:
                print(f"  {i:2d}. ❌ FILE NOT FOUND: {file_path}")
        
        total_size_str = self._format_file_size(total_size)
        print(f"\n📊 Total: {len(self.selected_files)} files, {total_size_str}")
    
    def do_clear_selection(self, arg):
        """Clear all selected files"""
        count = len(self.selected_files)
        self.selected_files = []
        print(f"🗑️  Cleared {count} selected files")
    
    def do_convert(self, arg):
        """Convert selected files: convert <operation> [output_dir]"""
        if not self.selected_files:
            print("❌ No files selected. Use 'select' command first.")
            return
        
        if not arg:
            print("❌ Usage: convert <operation> [output_dir]")
            print("   Available operations: pdf2docx, docx2pdf, csv2xlsx")
            print("   Example: convert pdf2docx ./converted")
            return
        
        args = shlex.split(arg)
        operation = args[0]
        output_dir = args[1] if len(args) > 1 else "./converted"
        
        # Validate operation
        supported_operations = {
            "pdf2docx": "PDF to DOCX",
            "docx2pdf": "DOCX to PDF", 
            "csv2xlsx": "CSV to XLSX"
        }
        
        if operation not in supported_operations:
            print(f"❌ Unsupported operation: {operation}")
            print(f"   Available: {', '.join(supported_operations.keys())}")
            return
        
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        print(f"🔄 Converting {len(self.selected_files)} files: {supported_operations[operation]}")
        print(f"📁 Output directory: {output_dir}")
        print("-" * 50)
        
        successful = 0
        failed = 0
        
        for file_path in self.selected_files:
            if not os.path.exists(file_path):
                print(f"❌ File not found: {file_path}")
                failed += 1
                continue
            
            try:
                result = self._convert_single_file(file_path, operation, output_dir)
                if result.get('success'):
                    print(f"✅ {os.path.basename(file_path)} -> {os.path.basename(result['output_path'])}")
                    successful += 1
                else:
                    print(f"❌ {os.path.basename(file_path)}: {result.get('error', 'Unknown error')}")
                    failed += 1
            except Exception as e:
                print(f"❌ {os.path.basename(file_path)}: {str(e)}")
                failed += 1
        
        print("-" * 50)
        print(f"📊 Conversion complete: {successful} successful, {failed} failed")
        
        if successful > 0:
            print(f"💾 Output files saved to: {output_dir}")
    
    def _convert_single_file(self, input_path: str, operation: str, output_dir: str) -> Dict[str, Any]:
        """Convert a single file"""
        input_name = os.path.basename(input_path)
        input_stem = os.path.splitext(input_name)[0]
        
        # Determine output extension and path
        output_extensions = {
            "pdf2docx": ".docx",
            "docx2pdf": ".pdf", 
            "csv2xlsx": ".xlsx"
        }
        
        output_ext = output_extensions.get(operation)
        output_path = os.path.join(output_dir, f"{input_stem}{output_ext}")
        
        # Avoid overwriting
        counter = 1
        original_output_path = output_path
        while os.path.exists(output_path):
            output_path = os.path.join(output_dir, f"{input_stem}_{counter}{output_ext}")
            counter += 1
        
        try:
            # Perform conversion based on operation
            if operation == "pdf2docx":
                result = self._convert_pdf_to_docx(input_path, output_path)
            elif operation == "docx2pdf":
                result = self._convert_docx_to_pdf(input_path, output_path)
            elif operation == "csv2xlsx":
                result = self._convert_csv_to_xlsx(input_path, output_path)
            else:
                result = {"success": False, "error": f"Unsupported operation: {operation}"}
            
            if result.get('success'):
                result['output_path'] = output_path
            return result
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _convert_pdf_to_docx(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert PDF to DOCX"""
        try:
            # Simple text extraction and DOCX creation as fallback
            import PyPDF2
            from docx import Document
            
            doc = Document()
            with open(input_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        doc.add_paragraph(text)
            
            doc.save(output_path)
            return {"success": True, "message": "PDF converted to DOCX (text only)"}
            
        except ImportError:
            return {"success": False, "error": "Required libraries not installed: PyPDF2, python-docx"}
        except Exception as e:
            return {"success": False, "error": f"Conversion failed: {str(e)}"}
    
    def _convert_docx_to_pdf(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert DOCX to PDF"""
        try:
            # Simple conversion using reportlab
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            
            doc = Document(input_path)
            c = canvas.Canvas(output_path, pagesize=letter)
            y_position = 750  # Start from top
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    if y_position < 50:  # New page if near bottom
                        c.showPage()
                        y_position = 750
                    
                    c.drawString(50, y_position, text[:80])  # Simple text placement
                    y_position -= 20
            
            c.save()
            return {"success": True, "message": "DOCX converted to PDF (basic formatting)"}
            
        except ImportError:
            return {"success": False, "error": "Required libraries not installed: python-docx, reportlab"}
        except Exception as e:
            return {"success": False, "error": f"Conversion failed: {str(e)}"}
    
    def _convert_csv_to_xlsx(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert CSV to XLSX"""
        try:
            import pandas as pd
            
            df = pd.read_csv(input_path)
            df.to_excel(output_path, index=False, engine='openpyxl')
            return {"success": True, "message": "CSV converted to XLSX"}
            
        except ImportError:
            return {"success": False, "error": "Required libraries not installed: pandas, openpyxl"}
        except Exception as e:
            return {"success": False, "error": f"Conversion failed: {str(e)}"}
    
    def _format_file_size(self, size_bytes):
        """Format file size in human-readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"
    
    def do_list(self, arg):
        """List available conversion tools"""
        tools = [
            {"name": "PDF to DOCX", "status": "✅ Basic", "formats": "PDF → DOCX"},
            {"name": "DOCX to PDF", "status": "✅ Basic", "formats": "DOCX → PDF"},
            {"name": "CSV to XLSX", "status": "✅ Full", "formats": "CSV → XLSX"},
            {"name": "Image Format Conversion", "status": "❌ Not available", "formats": "Install PIL"},
            {"name": "TXT to PDF", "status": "❌ Not available", "formats": "Install reportlab"},
        ]
        
        print("\n🛠️  CONVERSION TOOLS:")
        print("=" * 60)
        for tool in tools:
            print(f"  {tool['status']} {tool['name']:25} - {tool['formats']}")
        print("\n💡 Basic = text only, Full = formatting preserved")
    
    def do_status(self, arg):
        """Show system status"""
        # Check for required libraries
        libraries = {
            "pandas": "CSV to XLSX",
            "openpyxl": "CSV to XLSX", 
            "PyPDF2": "PDF processing",
            "python-docx": "DOCX processing",
            "reportlab": "PDF generation"
        }
        
        available_libs = []
        missing_libs = []
        
        for lib, purpose in libraries.items():
            try:
                __import__(lib)
                available_libs.append((lib, purpose, "✅"))
            except ImportError:
                missing_libs.append((lib, purpose, "❌"))
        
        print("\n📊 SYSTEM STATUS:")
        print("=" * 50)
        print(f"📁 Current directory: {self.current_directory}")
        print(f"📋 Selected files: {len(self.selected_files)}")
        print(f"💾 Available libraries:")
        
        for lib, purpose, status in available_libs + missing_libs:
            print(f"  {status} {lib:15} - {purpose}")
        
        if missing_libs:
            print(f"\n💡 Install missing libraries: pip install {' '.join([lib for lib, _, _ in missing_libs])}")
    
    def do_clear(self, arg):
        """Clear the screen"""
        os.system('cls' if os.name == 'nt' else 'clear')
        print(self.intro)
    
    def do_exit(self, arg):
        """Exit the interactive shell"""
        print("👋 Goodbye!")
        return True
    
    # Aliases
    def do_quit(self, arg):
        return self.do_exit(arg)
    
    def do_q(self, arg):
        return self.do_exit(arg)
    
    def emptyline(self):
        pass
    
    def default(self, line):
        print(f"❌ Unknown command: {line}")
        print("   Type 'help' for available commands")


def start_interactive_mode():
    """Start the interactive shell"""
    try:
        shell = InteractiveMode()
        shell.cmdloop()
        return {"success": True, "message": "Interactive session completed"}
    except KeyboardInterrupt:
        print("\n👋 Goodbye!")
        return {"success": True, "message": "Session interrupted by user"}
    except Exception as e:
        return {"success": False, "error": f"Interactive mode error: {e}"}