"""
ULTIMATE PDF to DOCX Converter - Professional Quality
Multi-library approach with intelligent fallback and quality optimization
Enhanced with async support, progress tracking, and MCP tool integration
"""

import os
import sys
import logging
import asyncio
from typing import List, Dict, Optional, Any, Tuple, Callable
from pathlib import Path
from datetime import datetime
import tempfile
import shutil
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from enum import Enum

# Add project root to path for imports
project_root = Path(__file__).parent.parent.parent.parent.parent
sys.path.insert(0, str(project_root))

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ConversionQuality(Enum):
    """Quality levels for conversion methods."""
    EXCELLENT = "excellent"
    VERY_GOOD = "very_good"
    GOOD = "good"
    BASIC = "basic"
    MINIMAL = "minimal"


class ConversionMethod(Enum):
    """Available conversion methods."""
    PDF2DOCX = "pdf2docx"
    PYMUPDF = "pymupdf"
    PDFPLUMBER = "pdfplumber"
    PYPDF = "pypdf"
    BASIC = "basic"


@dataclass
class ConversionResult:
    """Result of a PDF to DOCX conversion."""
    success: bool
    method: Optional[str] = None
    quality: Optional[ConversionQuality] = None
    output_path: Optional[str] = None
    error: Optional[str] = None
    features: List[str] = field(default_factory=list)
    stats: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "success": self.success,
            "method": self.method,
            "quality": self.quality.value if self.quality else None,
            "output_path": self.output_path,
            "error": self.error,
            "features": self.features,
            "stats": self.stats,
            "warnings": self.warnings
        }


@dataclass
class ConversionConfig:
    """Configuration for PDF conversion."""
    enable_ocr: bool = False
    extract_images: bool = True
    detect_headers_footers: bool = True
    extract_tables: bool = True
    preserve_formatting: bool = True
    multi_processing: bool = True
    max_image_width: float = 6.0  # inches
    header_footer_threshold: float = 0.15  # 15% of page
    
    # Quality thresholds
    min_quality_level: ConversionQuality = ConversionQuality.GOOD
    
    # Performance settings
    max_workers: int = 4
    timeout_seconds: int = 300


class ConversionMethodHandler:
    """Base class for conversion method handlers."""
    
    def __init__(self, config: ConversionConfig):
        self.config = config
        self.temp_dir = None
        
    def is_available(self) -> bool:
        """Check if this method is available."""
        raise NotImplementedError
    
    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        """Perform the conversion."""
        raise NotImplementedError
    
    def get_quality(self) -> ConversionQuality:
        """Get the quality level of this method."""
        raise NotImplementedError


class PDF2DOCXHandler(ConversionMethodHandler):
    """Handler for pdf2docx library (best quality)."""
    
    def is_available(self) -> bool:
        try:
            from pdf2docx import Converter
            return True
        except ImportError:
            return False
    
    def get_quality(self) -> ConversionQuality:
        return ConversionQuality.EXCELLENT
    
    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        try:
            from pdf2docx import Converter
            
            cv = Converter(input_path)
            cv.convert(
                output_path, 
                start=0, 
                end=None,
                multi_processing=self.config.multi_processing
            )
            cv.close()
            
            return ConversionResult(
                success=True,
                method=ConversionMethod.PDF2DOCX.value,
                quality=self.get_quality(),
                output_path=output_path,
                features=[
                    "layout_preservation", 
                    "table_extraction", 
                    "image_extraction",
                    "formatting_preservation", 
                    "headers_footers", 
                    "multi_column"
                ]
            )
        except Exception as e:
            logger.error(f"PDF2DOCX conversion failed: {e}")
            return ConversionResult(
                success=False,
                error=f"pdf2docx failed: {str(e)}"
            )


class PyMuPDFHandler(ConversionMethodHandler):
    """Handler for PyMuPDF/fitz library."""
    
    def __init__(self, config: ConversionConfig):
        super().__init__(config)
        self.images_extracted = 0
        self.header_footer_count = 0
    
    def is_available(self) -> bool:
        try:
            import fitz
            from docx import Document
            return True
        except ImportError:
            return False
    
    def get_quality(self) -> ConversionQuality:
        return ConversionQuality.VERY_GOOD
    
    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        try:
            import fitz
            from docx import Document
            from docx.shared import Inches, Pt
            from PIL import Image
            import io
            
            # Create temp directory for images
            self.temp_dir = tempfile.mkdtemp()
            
            doc = Document()
            pdf_doc = fitz.open(input_path)
            
            # Set document properties
            doc.core_properties.title = os.path.basename(input_path)
            doc.core_properties.author = "PDF Converter Pro"
            doc.core_properties.created = datetime.now()
            
            # Analyze page structure if header/footer detection is enabled
            page_structure = None
            if self.config.detect_headers_footers:
                page_structure = self._analyze_page_structure(pdf_doc)
            
            # Process each page
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                
                # Detect headers and footers
                header_footer_regions = None
                if self.config.detect_headers_footers and page_structure:
                    header_footer_regions = self._detect_header_footer(page, page_structure)
                
                # Extract text with formatting
                text_blocks = page.get_text("dict")
                blocks_sorted = self._sort_blocks_by_reading_order(text_blocks["blocks"])
                
                for block in blocks_sorted:
                    # Skip header/footer regions if enabled
                    if (header_footer_regions and 
                        self._is_in_header_footer_region(block, header_footer_regions)):
                        continue
                    
                    # Process text blocks
                    if "lines" in block:
                        self._add_text_block_to_docx(doc, block)
                    
                    # Process images if enabled
                    elif "image" in block and self.config.extract_images:
                        self._extract_and_add_image(doc, block, page_num)
                
                # Add page break (except last page)
                if page_num < len(pdf_doc) - 1:
                    doc.add_page_break()
            
            pdf_doc.close()
            doc.save(output_path)
            
            # Cleanup temp files
            self._cleanup_temp_files()
            
            return ConversionResult(
                success=True,
                method=ConversionMethod.PYMUPDF.value,
                quality=self.get_quality(),
                output_path=output_path,
                features=[
                    "text_formatting", 
                    "font_preservation", 
                    "layout_basic",
                    "image_extraction", 
                    "header_footer_detection"
                ],
                stats={
                    "images_extracted": self.images_extracted,
                    "headers_footers_detected": self.header_footer_count
                }
            )
            
        except Exception as e:
            logger.error(f"PyMuPDF conversion failed: {e}")
            self._cleanup_temp_files()
            return ConversionResult(
                success=False,
                error=f"PyMuPDF failed: {str(e)}"
            )
    
    def _analyze_page_structure(self, pdf_doc) -> Dict[str, Any]:
        """Analyze PDF structure for header/footer detection."""
        structure = {
            'page_heights': [],
            'common_top_regions': [],
            'common_bottom_regions': []
        }
        
        sample_pages = min(5, len(pdf_doc))
        
        for page_num in range(sample_pages):
            page = pdf_doc[page_num]
            structure['page_heights'].append(page.rect.height)
            
            text_blocks = page.get_text("dict")["blocks"]
            for block in text_blocks:
                if block['bbox'][1] < 100:
                    structure['common_top_regions'].append(block['bbox'][1])
                if block['bbox'][3] > page.rect.height - 100:
                    structure['common_bottom_regions'].append(block['bbox'][3])
        
        return structure
    
    def _detect_header_footer(self, page, page_structure) -> Dict[str, float]:
        """Detect header and footer regions."""
        page_height = page.rect.height
        threshold = self.config.header_footer_threshold
        
        regions = {
            'header_top': 0,
            'header_bottom': page_height * threshold,
            'footer_top': page_height * (1 - threshold),
            'footer_bottom': page_height
        }
        
        # Count header/footer content
        text_blocks = page.get_text("dict")["blocks"]
        header_footer_found = False
        
        for block in text_blocks:
            if "lines" in block:
                block_top = block['bbox'][1]
                block_bottom = block['bbox'][3]
                
                if (block_bottom < regions['header_bottom'] or 
                    block_top > regions['footer_top']):
                    header_footer_found = True
                    break
        
        if header_footer_found:
            self.header_footer_count += 1
        
        return regions
    
    def _is_in_header_footer_region(self, block, regions) -> bool:
        """Check if block is in header/footer region."""
        if 'bbox' not in block:
            return False
            
        block_top = block['bbox'][1]
        block_bottom = block['bbox'][3]
        
        return (block_bottom <= regions['header_bottom'] or 
                block_top >= regions['footer_top'])
    
    def _sort_blocks_by_reading_order(self, blocks: List[Dict]) -> List[Dict]:
        """Sort blocks by reading order (top-to-bottom, left-to-right)."""
        return sorted(blocks, key=lambda x: (x.get('bbox', [0, 0])[1], x.get('bbox', [0, 0])[0]))
    
    def _add_text_block_to_docx(self, doc, block):
        """Add text block with formatting to DOCX."""
        from docx.shared import Pt
        
        for line in block.get("lines", []):
            para = doc.add_paragraph()
            
            for span in line.get("spans", []):
                text = span.get("text", "").strip()
                if not text:
                    continue
                
                run = para.add_run(text + " ")
                
                # Apply formatting
                font_size = span.get("size", 11)
                run.font.size = Pt(min(font_size, 36))
                
                # Font styling
                font_name = span.get("font", "").lower()
                if "bold" in font_name or font_size > 12:
                    run.bold = True
                if "italic" in font_name:
                    run.italic = True
    
    def _extract_and_add_image(self, doc, block, page_num: int):
        """Extract and add image to DOCX."""
        try:
            from docx.shared import Inches
            from PIL import Image
            
            image_data = block.get("image")
            if not image_data:
                return
            
            image_ext = block.get("ext", "png")
            temp_image_path = os.path.join(
                self.temp_dir, 
                f"image_p{page_num}_{self.images_extracted}.{image_ext}"
            )
            
            with open(temp_image_path, "wb") as f:
                f.write(image_data)
            
            # Add image with size optimization
            with Image.open(temp_image_path) as img:
                width, height = img.size
                max_width_px = self.config.max_image_width * 72
                
                if width > max_width_px:
                    ratio = max_width_px / width
                    new_width = self.config.max_image_width
                    new_height = (height * ratio) / 72
                else:
                    new_width = width / 72
                    new_height = height / 72
                
                doc.add_picture(
                    temp_image_path, 
                    width=Inches(min(new_width, self.config.max_image_width))
                )
                self.images_extracted += 1
                
        except Exception as e:
            logger.warning(f"Image extraction failed: {e}")
    
    def _cleanup_temp_files(self):
        """Clean up temporary files."""
        try:
            if self.temp_dir and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception as e:
            logger.warning(f"Temp cleanup failed: {e}")


class PDFPlumberHandler(ConversionMethodHandler):
    """Handler for pdfplumber library (excellent for tables)."""
    
    def __init__(self, config: ConversionConfig):
        super().__init__(config)
        self.tables_extracted = 0
    
    def is_available(self) -> bool:
        try:
            import pdfplumber
            from docx import Document
            return True
        except ImportError:
            return False
    
    def get_quality(self) -> ConversionQuality:
        return ConversionQuality.GOOD
    
    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        try:
            import pdfplumber
            from docx import Document
            
            doc = Document()
            pdf = pdfplumber.open(input_path)
            
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract tables if enabled
                if self.config.extract_tables:
                    tables = self._extract_tables_enhanced(page)
                    for table_num, table_data in enumerate(tables, 1):
                        if table_data and any(any(cell for cell in row) for row in table_data):
                            self._add_table_to_docx(doc, table_data, f"Table {table_num}")
                
                # Extract text
                text = page.extract_text()
                if text and text.strip():
                    self._add_structured_text(doc, text)
            
            pdf.close()
            doc.save(output_path)
            
            return ConversionResult(
                success=True,
                method=ConversionMethod.PDFPLUMBER.value,
                quality=self.get_quality(),
                output_path=output_path,
                features=[
                    "table_extraction", 
                    "text_extraction", 
                    "structured_content"
                ],
                stats={"tables_extracted": self.tables_extracted}
            )
            
        except Exception as e:
            logger.error(f"PDFPlumber conversion failed: {e}")
            return ConversionResult(
                success=False,
                error=f"pdfplumber failed: {str(e)}"
            )
    
    def _extract_tables_enhanced(self, page) -> List[List[List[str]]]:
        """Enhanced table extraction."""
        try:
            tables = page.extract_tables({
                "vertical_strategy": "lines",
                "horizontal_strategy": "lines",
                "snap_tolerance": 3,
                "join_tolerance": 3,
                "edge_min_length": 3
            })
            
            enhanced_tables = []
            for table in tables:
                if table and any(any(cell for cell in row) for row in table):
                    cleaned_table = self._clean_table_data(table)
                    if cleaned_table:
                        enhanced_tables.append(cleaned_table)
                        self.tables_extracted += 1
            
            return enhanced_tables
            
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")
            return []
    
    def _clean_table_data(self, table_data: List[List[str]]) -> List[List[str]]:
        """Clean and normalize table data."""
        if not table_data:
            return []
        
        # Remove empty rows
        cleaned_data = [
            row for row in table_data 
            if any(cell and str(cell).strip() for cell in row)
        ]
        
        if not cleaned_data:
            return []
        
        # Normalize column count
        max_cols = max(len(row) for row in cleaned_data)
        normalized_data = []
        
        for row in cleaned_data:
            normalized_row = list(row) + [''] * (max_cols - len(row))
            cleaned_row = [
                str(cell).strip() if cell else '' 
                for cell in normalized_row
            ]
            normalized_data.append(cleaned_row)
        
        return normalized_data
    
    def _add_table_to_docx(self, doc, table_data, title=None):
        """Add table to DOCX with formatting."""
        if title:
            doc.add_paragraph(title, style='Heading 3')
        
        if not table_data:
            return
        
        table = doc.add_table(
            rows=len(table_data), 
            cols=max(len(row) for row in table_data)
        )
        table.style = 'Table Grid'
        
        for row_idx, row in enumerate(table_data):
            for col_idx, cell in enumerate(row):
                if cell and col_idx < len(table.rows[row_idx].cells):
                    table_cell = table.cell(row_idx, col_idx)
                    table_cell.text = str(cell)
                    
                    # Bold header row
                    if row_idx == 0:
                        for paragraph in table_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    def _add_structured_text(self, doc, text):
        """Add text with basic structure detection."""
        paragraphs = text.split('\n\n')
        
        for para_text in paragraphs:
            clean_text = para_text.strip()
            if clean_text:
                # Simple heading detection
                if self._is_heading(clean_text):
                    doc.add_heading(clean_text, level=2)
                else:
                    doc.add_paragraph(clean_text)
    
    def _is_heading(self, text: str) -> bool:
        """Detect if text is a heading."""
        return (
            len(text) < 100 and
            (text.isupper() or text.endswith(':')) and
            not text.endswith('.') and
            not text.endswith(',')
        )


class PyPDFHandler(ConversionMethodHandler):
    """Handler for PyPDF library (basic fallback)."""
    
    def is_available(self) -> bool:
        try:
            from pypdf import PdfReader
            from docx import Document
            return True
        except ImportError:
            return False
    
    def get_quality(self) -> ConversionQuality:
        return ConversionQuality.BASIC
    
    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        try:
            from pypdf import PdfReader
            from docx import Document
            
            doc = Document()
            reader = PdfReader(input_path)
            
            # Add metadata
            if reader.metadata:
                if reader.metadata.title:
                    doc.add_heading(reader.metadata.title, 0)
                if reader.metadata.author:
                    doc.add_paragraph(f"Author: {reader.metadata.author}")
                doc.add_paragraph("---")
            
            # Extract text
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    for line in text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
            
            doc.save(output_path)
            
            return ConversionResult(
                success=True,
                method=ConversionMethod.PYPDF.value,
                quality=self.get_quality(),
                output_path=output_path,
                features=["text_extraction", "metadata"],
                warnings=["Limited formatting preservation"]
            )
            
        except Exception as e:
            logger.error(f"PyPDF conversion failed: {e}")
            return ConversionResult(
                success=False,
                error=f"pypdf failed: {str(e)}"
            )


class UltimatePDFToDOCXConverter:
    """
    Ultimate PDF to DOCX converter with intelligent method selection.
    Tries multiple libraries in order of quality until successful.
    """
    
    def __init__(self, config: Optional[ConversionConfig] = None):
        self.config = config or ConversionConfig()
        self.handlers = self._initialize_handlers()
        
    def _initialize_handlers(self) -> Dict[ConversionMethod, ConversionMethodHandler]:
        """Initialize all available conversion handlers."""
        handlers = {
            ConversionMethod.PDF2DOCX: PDF2DOCXHandler(self.config),
            ConversionMethod.PYMUPDF: PyMuPDFHandler(self.config),
            ConversionMethod.PDFPLUMBER: PDFPlumberHandler(self.config),
            ConversionMethod.PYPDF: PyPDFHandler(self.config),
        }
        return handlers
    
    def get_available_methods(self) -> List[Tuple[ConversionMethod, ConversionQuality]]:
        """Get list of available conversion methods with their quality levels."""
        available = []
        for method, handler in self.handlers.items():
            if handler.is_available():
                available.append((method, handler.get_quality()))
        return available
    
    def convert(self,
                input_path: str,
                output_path: Optional[str] = None,
                progress_callback: Optional[Callable[[str, float], None]] = None
                ) -> ConversionResult:
        """
        Convert PDF to DOCX using the best available method.
        
        Args:
            input_path: Path to input PDF file
            output_path: Output DOCX path (auto-generated if None)
            progress_callback: Optional callback for progress updates (message, percentage)
            
        Returns:
            ConversionResult with detailed information
        """
        # Validate input
        if not os.path.exists(input_path):
            return ConversionResult(
                success=False,
                error=f"Input file not found: {input_path}"
            )
        
        if not input_path.lower().endswith('.pdf'):
            return ConversionResult(
                success=False,
                error="Input file must be a PDF"
            )
        
        # Generate output path if not provided
        if output_path is None:
            output_path = self._generate_output_path(input_path)
        
        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Try conversion methods in order of quality
        available_methods = self.get_available_methods()
        if not available_methods:
            return ConversionResult(
                success=False,
                error="No conversion methods available. Install dependencies: "
                      "pip install pdf2docx pymupdf pdfplumber pypdf python-docx"
            )
        
        attempts = []
        start_time = datetime.now()
        
        for idx, (method, quality) in enumerate(available_methods):
            if progress_callback:
                progress = (idx / len(available_methods)) * 100
                progress_callback(f"Trying {method.value}...", progress)
            
            logger.info(f"Attempting conversion with: {method.value} (quality: {quality.value})")
            
            handler = self.handlers[method]
            result = handler.convert(input_path, output_path)
            
            if result.success:
                # Add conversion statistics
                result.stats.update(self._calculate_stats(input_path, output_path, start_time))
                
                if progress_callback:
                    progress_callback("Conversion complete!", 100)
                
                logger.info(f"✓ Conversion successful using {method.value}")
                return result
            else:
                attempts.append(f"{method.value}: {result.error}")
                logger.warning(f"✗ {method.value} failed: {result.error}")
        
        # All methods failed
        return ConversionResult(
            success=False,
            error="All conversion methods failed",
            warnings=attempts
        )
    
    async def convert_async(self,
                           input_path: str,
                           output_path: Optional[str] = None,
                           progress_callback: Optional[Callable[[str, float], None]] = None
                           ) -> ConversionResult:
        """Async version of convert method."""
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor(max_workers=self.config.max_workers) as executor:
            result = await loop.run_in_executor(
                executor,
                self.convert,
                input_path,
                output_path,
                progress_callback
            )
        return result
    
    def batch_convert(self,
                     input_files: List[str],
                     output_dir: Optional[str] = None,
                     progress_callback: Optional[Callable[[str, float], None]] = None
                     ) -> List[ConversionResult]:
        """
        Convert multiple PDF files to DOCX.
        
        Args:
            input_files: List of input PDF file paths
            output_dir: Output directory (uses input dir if None)
            progress_callback: Optional progress callback
            
        Returns:
            List of ConversionResults
        """
        results = []
        total_files = len(input_files)
        
        for idx, input_file in enumerate(input_files):
            if progress_callback:
                progress = (idx / total_files) * 100
                progress_callback(f"Converting {os.path.basename(input_file)}", progress)
            
            output_path = None
            if output_dir:
                filename = Path(input_file).stem + ".docx"
                output_path = os.path.join(output_dir, filename)
            
            result = self.convert(input_file, output_path)
            results.append(result)
        
        if progress_callback:
            progress_callback("Batch conversion complete!", 100)
        
        return results
    
    def _generate_output_path(self, input_path: str) -> str:
        """Generate unique output path."""
        base_name = os.path.splitext(input_path)[0]
        output_path = f"{base_name}.docx"
        
        counter = 1
        while os.path.exists(output_path):
            output_path = f"{base_name}_{counter}.docx"
            counter += 1
        
        return output_path
    
    def _calculate_stats(self, input_path: str, output_path: str, start_time: datetime) -> Dict[str, Any]:
        """Calculate conversion statistics."""
        end_time = datetime.now()
        duration = (end_time - start_time).total_seconds()
        
        input_size = os.path.getsize(input_path)
        output_size = os.path.getsize(output_path)
        
        return {
            "input_file": os.path.basename(input_path),
            "output_file": os.path.basename(output_path),
            "input_size": self._format_size(input_size),
            "output_size": self._format_size(output_size),
            "size_ratio": f"{(output_size/input_size*100):.1f}%" if input_size > 0 else "N/A",
            "conversion_time": f"{duration:.2f}s",
            "timestamp": end_time.isoformat(),
            "output_path": os.path.abspath(output_path)
        }
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human-readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"
    
    def get_system_info(self) -> Dict[str, Any]:
        """Get system information about available methods."""
        info = {}
        for method, handler in self.handlers.items():
            info[method.value] = {
                "available": handler.is_available(),
                "quality": handler.get_quality().value if handler.is_available() else "N/A"
            }
        return info


# MCP Tool Integration
try:
    from ai_infrastructure.mcp.tool_registry import BaseTool
    
    class PDFToDOCXConverter(BaseTool):
        """MCP Tool for PDF to DOCX conversion."""
        
        def __init__(self):
            super().__init__(
                name="convert_pdf_to_docx",
                description="Convert PDF files to DOCX format with professional quality using multiple conversion methods",
                parameters={
                    "input_path": {
                        "type": "str",
                        "required": True,
                        "description": "Path to the input PDF file"
                    },
                    "output_path": {
                        "type": "str",
                        "required": False,
                        "description": "Output DOCX file path (auto-generated if not provided)"
                    },
                    "extract_images": {
                        "type": "bool",
                        "required": False,
                        "description": "Extract and include images from PDF",
                        "default": True
                    },
                    "detect_headers_footers": {
                        "type": "bool",
                        "required": False,
                        "description": "Detect and remove headers/footers",
                        "default": True
                    },
                    "extract_tables": {
                        "type": "bool",
                        "required": False,
                        "description": "Extract and format tables",
                        "default": True
                    }
                }
            )
            self.converter = None
        
        def run(self, 
                input_path: str,
                output_path: Optional[str] = None,
                extract_images: bool = True,
                detect_headers_footers: bool = True,
                extract_tables: bool = True,
                **kwargs) -> Dict[str, Any]:
            """
            Execute PDF to DOCX conversion.
            
            Args:
                input_path: Path to input PDF file
                output_path: Output DOCX path (optional)
                extract_images: Whether to extract images
                detect_headers_footers: Whether to detect headers/footers
                extract_tables: Whether to extract tables
                
            Returns:
                Conversion result dictionary
            """
            # Create config
            config = ConversionConfig(
                extract_images=extract_images,
                detect_headers_footers=detect_headers_footers,
                extract_tables=extract_tables
            )
            
            # Create converter
            self.converter = UltimatePDFToDOCXConverter(config)
            
            # Perform conversion
            result = self.converter.convert(input_path, output_path)
            
            return result.to_dict()
    
except ImportError:
    logger.warning("MCP tool_registry not available - MCP integration disabled")
    PDFToDOCXConverter = None


# Convenience functions for direct usage
def convert_pdf_to_docx(
    input_path: str,
    output_path: Optional[str] = None,
    extract_images: bool = True,
    detect_headers_footers: bool = True,
    extract_tables: bool = True,
    progress_callback: Optional[Callable[[str, float], None]] = None
) -> Dict[str, Any]:
    """
    Convert PDF to DOCX with professional quality.
    
    Args:
        input_path: Path to input PDF file
        output_path: Output DOCX path (auto-generated if None)
        extract_images: Extract and include images
        detect_headers_footers: Detect and remove headers/footers
        extract_tables: Extract and format tables
        progress_callback: Optional progress callback function
        
    Returns:
        Dictionary with conversion results
        
    Example:
        >>> result = convert_pdf_to_docx("document.pdf", "output.docx")
        >>> if result["success"]:
        >>>     print(f"Converted using {result['method']}")
        >>>     print(f"Output: {result['output_path']}")
    """
    config = ConversionConfig(
        extract_images=extract_images,
        detect_headers_footers=detect_headers_footers,
        extract_tables=extract_tables
    )
    
    converter = UltimatePDFToDOCXConverter(config)
    result = converter.convert(input_path, output_path, progress_callback)
    
    return result.to_dict()


async def convert_pdf_to_docx_async(
    input_path: str,
    output_path: Optional[str] = None,
    **kwargs
) -> Dict[str, Any]:
    """
    Async version of PDF to DOCX conversion.
    
    Args:
        input_path: Path to input PDF file
        output_path: Output DOCX path (auto-generated if None)
        **kwargs: Additional configuration options
        
    Returns:
        Dictionary with conversion results
    """
    config = ConversionConfig(**{k: v for k, v in kwargs.items() if k in ConversionConfig.__annotations__})
    converter = UltimatePDFToDOCXConverter(config)
    result = await converter.convert_async(input_path, output_path)
    return result.to_dict()


def batch_convert_pdfs(
    input_files: List[str],
    output_dir: Optional[str] = None,
    progress_callback: Optional[Callable[[str, float], None]] = None,
    **kwargs
) -> List[Dict[str, Any]]:
    """
    Convert multiple PDF files to DOCX.
    
    Args:
        input_files: List of input PDF file paths
        output_dir: Output directory (uses input dir if None)
        progress_callback: Optional progress callback
        **kwargs: Additional configuration options
        
    Returns:
        List of conversion result dictionaries
        
    Example:
        >>> files = ["doc1.pdf", "doc2.pdf", "doc3.pdf"]
        >>> results = batch_convert_pdfs(files, "output_folder")
        >>> successful = sum(1 for r in results if r["success"])
        >>> print(f"Converted {successful}/{len(files)} files")
    """
    config = ConversionConfig(**{k: v for k, v in kwargs.items() if k in ConversionConfig.__annotations__})
    converter = UltimatePDFToDOCXConverter(config)
    results = converter.batch_convert(input_files, output_dir, progress_callback)
    return [r.to_dict() for r in results]


def get_system_info() -> Dict[str, Any]:
    """
    Get information about available conversion methods.
    
    Returns:
        Dictionary with method availability and quality levels
        
    Example:
        >>> info = get_system_info()
        >>> for method, details in info.items():
        >>>     if details["available"]:
        >>>         print(f"{method}: {details['quality']}")
    """
    converter = UltimatePDFToDOCXConverter()
    return converter.get_system_info()


# CLI interface for testing
def main():
    """CLI interface for testing the converter."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Ultimate PDF to DOCX Converter",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert single file
  python convert_pdf_to_docx.py input.pdf
  
  # Convert with custom output
  python convert_pdf_to_docx.py input.pdf -o output.docx
  
  # Convert multiple files
  python convert_pdf_to_docx.py file1.pdf file2.pdf file3.pdf -d output_folder
  
  # Show available methods
  python convert_pdf_to_docx.py --info
        """
    )
    
    parser.add_argument(
        "input_files",
        nargs="*",
        help="Input PDF file(s)"
    )
    parser.add_argument(
        "-o", "--output",
        help="Output DOCX file path (for single file conversion)"
    )
    parser.add_argument(
        "-d", "--output-dir",
        help="Output directory (for batch conversion)"
    )
    parser.add_argument(
        "--no-images",
        action="store_true",
        help="Don't extract images"
    )
    parser.add_argument(
        "--no-tables",
        action="store_true",
        help="Don't extract tables"
    )
    parser.add_argument(
        "--no-header-footer",
        action="store_true",
        help="Don't detect headers/footers"
    )
    parser.add_argument(
        "--info",
        action="store_true",
        help="Show available conversion methods"
    )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Verbose output"
    )
    
    args = parser.parse_args()
    
    # Set logging level
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Show system info
    if args.info:
        print("\n" + "="*70)
        print("AVAILABLE CONVERSION METHODS")
        print("="*70)
        info = get_system_info()
        for method, details in info.items():
            status = "✓" if details["available"] else "✗"
            print(f"{status} {method:15s} - Quality: {details['quality']}")
        print("="*70 + "\n")
        
        if not any(d["available"] for d in info.values()):
            print("⚠️  No conversion methods available!")
            print("Install dependencies: pip install pdf2docx pymupdf pdfplumber pypdf python-docx")
        return
    
    # Validate input
    if not args.input_files:
        parser.print_help()
        return
    
    # Progress callback
    def progress_callback(message: str, percentage: float):
        if args.verbose:
            print(f"[{percentage:5.1f}%] {message}")
    
    # Single file conversion
    if len(args.input_files) == 1 and not args.output_dir:
        input_file = args.input_files[0]
        output_file = args.output if args.output else None
        
        print(f"\n📄 Converting: {input_file}")
        
        result = convert_pdf_to_docx(
            input_file,
            output_file,
            extract_images=not args.no_images,
            detect_headers_footers=not args.no_header_footer,
            extract_tables=not args.no_tables,
            progress_callback=progress_callback if args.verbose else None
        )
        
        if result["success"]:
            print(f"✓ Success! Method: {result['method']} (Quality: {result['quality']})")
            print(f"  Output: {result['output_path']}")
            if result.get("stats"):
                stats = result["stats"]
                print(f"  Size: {stats.get('input_size')} → {stats.get('output_size')} ({stats.get('size_ratio')})")
                print(f"  Time: {stats.get('conversion_time')}")
            if result.get("features"):
                print(f"  Features: {', '.join(result['features'])}")
        else:
            print(f"✗ Failed: {result['error']}")
            if result.get("warnings"):
                print("  Attempts:")
                for warning in result["warnings"]:
                    print(f"    - {warning}")
    
    # Batch conversion
    else:
        output_dir = args.output_dir or os.path.dirname(args.input_files[0]) or "."
        
        print(f"\n📚 Batch converting {len(args.input_files)} files to: {output_dir}")
        
        results = batch_convert_pdfs(
            args.input_files,
            output_dir,
            progress_callback=progress_callback if args.verbose else None,
            extract_images=not args.no_images,
            detect_headers_footers=not args.no_header_footer,
            extract_tables=not args.no_tables
        )
        
        # Summary
        successful = sum(1 for r in results if r["success"])
        failed = len(results) - successful
        
        print(f"\n{'='*70}")
        print(f"BATCH CONVERSION SUMMARY")
        print(f"{'='*70}")
        print(f"Total:      {len(results)}")
        print(f"Successful: {successful}")
        print(f"Failed:     {failed}")
        print(f"{'='*70}\n")
        
        # Show details
        for idx, (input_file, result) in enumerate(zip(args.input_files, results), 1):
            status = "✓" if result["success"] else "✗"
            filename = os.path.basename(input_file)
            print(f"{status} [{idx}/{len(results)}] {filename}")
            
            if result["success"]:
                if args.verbose and result.get("stats"):
                    print(f"    Method: {result['method']} | Time: {result['stats'].get('conversion_time')}")
            else:
                print(f"    Error: {result['error']}")


if __name__ == "__main__":
    main()