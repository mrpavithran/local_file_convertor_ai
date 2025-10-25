"""
Document ingestor for PDFs, DOCX, TXT, CSV with enhanced text processing.
Dependencies: pdfplumber, python-docx, pandas, beautifulsoup4 (for HTML)
"""
from typing import List, Dict, Optional, Callable
import os
import re
import json
import logging
from pathlib import Path
from tqdm import tqdm

import pdfplumber
import docx
import pandas as pd
from bs4 import BeautifulSoup

# Set up logging
logger = logging.getLogger(__name__)

class DocumentIngestor:
    def __init__(self, 
                 chunk_size: int = 800, 
                 overlap: int = 100,
                 min_chunk_size: int = 50,
                 smart_chunking: bool = True):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.min_chunk_size = min_chunk_size
        self.smart_chunking = smart_chunking
        
        # File extension to handler mapping
        self.reader_registry = {
            '.txt': self._read_txt,
            '.pdf': self._read_pdf,
            '.docx': self._read_docx,
            '.doc': self._read_docx,  # Handle .doc as .docx if python-docx supports it
            '.csv': self._read_csv,
            '.tsv': self._read_tsv,
            '.json': self._read_json,
            '.html': self._read_html,
            '.htm': self._read_html,
        }

    def _read_txt(self, path: str) -> str:
        """Read text file with encoding detection."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        # Fallback with error handling
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _read_pdf(self, path: str) -> str:
        """Read PDF with enhanced text extraction and layout preservation."""
        text = []
        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    # Try to extract text with layout preservation
                    page_text = page.extract_text(
                        layout=True,  # Preserve layout
                        x_tolerance=2,
                        y_tolerance=2
                    ) or ""
                    
                    # Fallback to simple extraction
                    if not page_text.strip():
                        page_text = page.extract_text() or ""
                    
                    # Extract tables if any
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = " | ".join([" | ".join(filter(None, row)) for row in table if any(row)])
                            if table_text:
                                page_text += f"\nTable: {table_text}"
                    
                    text.append(page_text)
        except Exception as e:
            logger.error(f"Error reading PDF {path}: {e}")
            return ""
        
        return "\n".join(text)

    def _read_docx(self, path: str) -> str:
        """Read DOCX file with formatting preservation."""
        try:
            doc = docx.Document(path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error reading DOCX {path}: {e}")
            return ""

    def _read_csv(self, path: str) -> str:
        """Read CSV file with pandas."""
        try:
            df = pd.read_csv(path, dtype=str, keep_default_na=False, encoding_errors='ignore')
            # Convert to readable text format
            rows = []
            # Add header
            rows.append(" | ".join(df.columns.tolist()))
            # Add rows
            for _, row in df.iterrows():
                rows.append(" | ".join(row.fillna('').astype(str).tolist()))
            return "\n".join(rows)
        except Exception as e:
            logger.error(f"Error reading CSV {path}: {e}")
            return ""

    def _read_tsv(self, path: str) -> str:
        """Read TSV file."""
        try:
            df = pd.read_csv(path, sep='\t', dtype=str, keep_default_na=False, encoding_errors='ignore')
            rows = []
            rows.append(" | ".join(df.columns.tolist()))
            for _, row in df.iterrows():
                rows.append(" | ".join(row.fillna('').astype(str).tolist()))
            return "\n".join(rows)
        except Exception as e:
            logger.error(f"Error reading TSV {path}: {e}")
            return ""

    def _read_json(self, path: str) -> str:
        """Read JSON file and convert to readable text."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            def flatten_json(obj, parent_key='', sep=': '):
                """Recursively flatten JSON object."""
                items = []
                if isinstance(obj, dict):
                    for k, v in obj.items():
                        new_key = f"{parent_key}{sep}{k}" if parent_key else k
                        items.extend(flatten_json(v, new_key, sep))
                elif isinstance(obj, list):
                    for i, v in enumerate(obj):
                        new_key = f"{parent_key}[{i}]"
                        items.extend(flatten_json(v, new_key, sep))
                else:
                    items.append(f"{parent_key}{sep}{obj}")
                return items
            
            flattened = flatten_json(data)
            return "\n".join(flattened)
        except Exception as e:
            logger.error(f"Error reading JSON {path}: {e}")
            return ""

    def _read_html(self, path: str) -> str:
        """Read HTML file and extract text content."""
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text and clean up
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            return '\n'.join(chunk for chunk in chunks if chunk)
        except Exception as e:
            logger.error(f"Error reading HTML {path}: {e}")
            return ""

    def clean_text(self, text: str) -> str:
        """Enhanced text cleaning with multiple normalization steps."""
        if not text:
            return ""
        
        # Normalize line endings
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        
        # Remove excessive whitespace
        text = re.sub(r'[ \t]{2,}', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Normalize Unicode characters (optional)
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII if desired
        
        # Remove unwanted characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-/\n]', '', text)
        
        return text.strip()

    def chunk_text(self, text: str, chunk_size: Optional[int] = None, overlap: Optional[int] = None) -> List[str]:
        """Enhanced chunking that respects sentence boundaries when possible."""
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.overlap
        
        text = text.strip()
        if not text or len(text) < self.min_chunk_size:
            return []
        
        # If text is smaller than chunk size, return as single chunk
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            # Calculate end position
            end = start + chunk_size
            
            # If this is not the first chunk, adjust for overlap
            if start > 0:
                start = max(0, start - overlap)
            
            # If we're at the end, take the remaining text
            if end >= text_length:
                chunk = text[start:].strip()
                if chunk and len(chunk) >= self.min_chunk_size:
                    chunks.append(chunk)
                break
            
            # Smart chunking: try to break at sentence boundaries
            if self.smart_chunking:
                # Look for sentence endings near the chunk boundary
                sentence_endings = ['.', '!', '?', '\n\n']
                for sep in sentence_endings:
                    # Look for separator near the end of chunk
                    pos = text.rfind(sep, start, end)
                    if pos != -1 and pos > start + chunk_size * 0.5:  # Only break if we're past halfway
                        end = pos + len(sep)
                        break
                else:
                    # No sentence ending found, try to break at word boundary
                    last_space = text.rfind(' ', start, end)
                    if last_space != -1 and last_space > start + chunk_size * 0.7:
                        end = last_space
            
            chunk = text[start:end].strip()
            if chunk and len(chunk) >= self.min_chunk_size:
                chunks.append(chunk)
            
            start = end
        
        return chunks

    def register_reader(self, extension: str, reader_func: Callable):
        """Register a custom reader function for a file extension."""
        self.reader_registry[extension.lower()] = reader_func

    def ingest_documents_from_paths(self, 
                                  paths: List[str], 
                                  chunk_size: Optional[int] = None,
                                  overlap: Optional[int] = None,
                                  show_progress: bool = True) -> List[Dict]:
        """Ingest documents from file paths with progress tracking."""
        docs = []
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.overlap
        
        iterable = paths
        if show_progress:
            iterable = tqdm(paths, desc="Ingesting documents")
        
        for path in iterable:
            if not os.path.exists(path):
                logger.warning(f"File not found: {path}")
                continue
            
            # Get file extension
            file_ext = Path(path).suffix.lower()
            
            # Get appropriate reader
            reader_func = self.reader_registry.get(file_ext, self._read_txt)
            
            try:
                text = reader_func(path)
                if not text.strip():
                    logger.warning(f"No content extracted from: {path}")
                    continue
                
                # Clean text
                cleaned_text = self.clean_text(text)
                if not cleaned_text:
                    continue
                
                # Chunk text
                chunks = self.chunk_text(cleaned_text, chunk_size, overlap)
                
                # Create document objects
                for i, chunk in enumerate(chunks):
                    docs.append({
                        "text": chunk,
                        "metadata": {
                            "source": os.path.abspath(path),
                            "file_name": os.path.basename(path),
                            "file_type": file_ext,
                            "chunk": i,
                            "total_chunks": len(chunks),
                            "chunk_size": len(chunk),
                            "word_count": len(chunk.split())
                        }
                    })
                
                logger.info(f"Processed {path}: {len(chunks)} chunks")
                
            except Exception as e:
                logger.error(f"Failed to process {path}: {e}")
                continue
        
        logger.info(f"Ingested {len(docs)} total chunks from {len(paths)} files")
        return docs

    def ingest_directory(self, 
                        directory: str, 
                        recursive: bool = True,
                        extensions: Optional[List[str]] = None,
                        **kwargs) -> List[Dict]:
        """Ingest all documents from a directory."""
        directory_path = Path(directory)
        if not directory_path.exists():
            raise ValueError(f"Directory does not exist: {directory}")
        
        # Find all files
        pattern = "**/*" if recursive else "*"
        all_files = [str(p) for p in directory_path.glob(pattern) if p.is_file()]
        
        # Filter by extensions if provided
        if extensions:
            extensions = [ext.lower() if ext.startswith('.') else f'.{ext.lower()}' for ext in extensions]
            all_files = [f for f in all_files if Path(f).suffix.lower() in extensions]
        
        return self.ingest_documents_from_paths(all_files, **kwargs)


# Convenience function for backward compatibility
def ingest_documents_from_paths(paths: List[str], chunk_size: int = 800, overlap: int = 100) -> List[Dict]:
    ingestor = DocumentIngestor(chunk_size=chunk_size, overlap=overlap)
    return ingestor.ingest_documents_from_paths(paths, show_progress=True)


# Example usage
if __name__ == "__main__":
    # Initialize ingestor
    ingestor = DocumentIngestor(chunk_size=1000, overlap=150, smart_chunking=True)
    
    # Ingest specific files
    documents = ingestor.ingest_documents_from_paths([
        "document.pdf",
        "data.csv", 
        "notes.txt"
    ])
    
    # Ingest entire directory
    documents = ingestor.ingest_directory(
        "./documents",
        extensions=[".pdf", ".docx", ".txt", ".csv"],
        recursive=True
    )
    
    print(f"Ingested {len(documents)} document chunks")